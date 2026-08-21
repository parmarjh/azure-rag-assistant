from __future__ import annotations

import re
from pathlib import Path

from .models import Document, Section

try:
    import pdfplumber
except ImportError:  # pragma: no cover
    pdfplumber = None
try:
    from docx import Document as DocxDocument
except ImportError:  # pragma: no cover
    DocxDocument = None
try:
    import openpyxl
except ImportError:  # pragma: no cover
    openpyxl = None

HEADING = re.compile(r"^\s*(\d+(?:\.\d+)*)\.?\s+(.+?)\s*$")
BOILERPLATE = re.compile(
    r"(?i)^\s*(?:[-—]\s*)?(?:internal use only|page\s+\d+|"
    r"\(cid:\d+\))\s*$"
)


def _clean_text(text: str) -> str:
    lines = []
    for raw_line in text.splitlines():
        line = re.sub(r"\(cid:\d+\)", "", raw_line, flags=re.I)
        line = re.sub(r"(?i)\s*[-—]\s*internal use only\b", "", line)
        line = re.sub(r"(?i)\bpage\s+\d+\b", "", line)
        line = re.sub(r"\s+", " ", line).strip()
        if not line or BOILERPLATE.match(line) or line.casefold() == "northwind traders, inc.":
            continue
        lines.append(line)
    return "\n".join(lines)


def _title(text: str, path: Path) -> str:
    first = next((x.strip() for x in text.splitlines() if x.strip()), path.stem)
    return re.sub(r"\s*\|\s*Northwind.*$", "", first).strip()


def _sections(text: str, pages: list[int] | None = None) -> list[Section]:
    lines = [re.sub(r"\s+", " ", x).strip() for x in text.splitlines()]
    sections: list[Section] = []
    current: Section | None = None
    for i, line in enumerate(lines):
        if not line:
            continue
        m = HEADING.match(line)
        table_row_heading = bool(
            re.match(r"^\d+\s*[–-]\s*\d+\b", line)
            and len(re.findall(r"\d+", line)) >= 3
        )
        if m and not table_row_heading and (len(line) < 140 or m.group(1).count(".") > 0):
            if current and current.text.strip():
                sections.append(current)
            current = Section(m.group(1) + " " + m.group(2), "", pages[i] if pages and i < len(pages) else None, m.group(1))
        elif current is None:
            current = Section("Document", line)
        else:
            current.text += ("\n" if current.text else "") + line
    if current and current.text.strip():
        sections.append(current)
    return sections or [Section("Document", text)]


def _xlsx_cell(value, header: str) -> str:
    if isinstance(value, (int, float)) and "%" in header and 0 <= value <= 1:
        return f"{value:.0%}"
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value) if value is not None else ""


def _xlsx_rows(sheet) -> list[str]:
    rows: list[str] = []
    headers: list[str] | None = None
    for row in sheet.iter_rows(values_only=True):
        values = list(row)
        if not any(value is not None for value in values):
            continue
        text_values = [str(value).strip() for value in values if value is not None]
        if (
            len(text_values) > 1
            and any("%" in value or "price" in value.lower() for value in text_values)
        ):
            headers = text_values
            rows.append("| " + " | ".join(text_values) + " |")
            continue
        if headers:
            fields = [
                f"{header}: {_xlsx_cell(value, header)}"
                for header, value in zip(headers, values)
                if value is not None
            ]
            rows.append("; ".join(fields))
        else:
            rows.append("| " + " | ".join(_xlsx_cell(value, "") for value in values) + " |")
    return rows


def parse_file(path: str | Path, department: str | None = None) -> Document:
    p = Path(path)
    department = department or p.parent.name
    text = ""
    pages: list[int] = []
    if p.suffix.lower() == ".pdf" and pdfplumber:
        with pdfplumber.open(p) as pdf:
            for page_no, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text() or ""
                text += page_text + "\n"
                pages.extend([page_no] * max(1, len(page_text.splitlines())))
    elif p.suffix.lower() == ".docx" and DocxDocument:
        d = DocxDocument(p)
        text = "\n".join(x.text for x in d.paragraphs if x.text.strip())
        for table in d.tables:
            text += "\n" + "\n".join(" | ".join(c.text.strip() for c in row.cells) for row in table.rows)
    elif p.suffix.lower() == ".xlsx" and openpyxl:
        book = openpyxl.load_workbook(p, data_only=True)
        blocks = []
        for sheet in book.worksheets:
            blocks.append(f"Sheet: {sheet.title}\n" + "\n".join(_xlsx_rows(sheet)))
        text = "\n\n".join(blocks)
    else:
        text = p.read_text(errors="replace")
    text = _clean_text(text)
    title = _title(text, p)
    doc_id = re.sub(r"[^a-z0-9]+", "-", p.stem.lower()).strip("-")
    return Document(doc_id, p.name, str(p), title, department, sections=_sections(text, pages))


def parse_directory(source_dir: str | Path) -> list[Document]:
    root = Path(source_dir)
    docs = [parse_file(p) for p in sorted(root.rglob("*")) if p.suffix.lower() in {".pdf", ".docx", ".xlsx"}]
    from .metadata import enrich_documents
    return enrich_documents(docs)
